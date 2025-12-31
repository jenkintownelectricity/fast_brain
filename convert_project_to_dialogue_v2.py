import os
import json
import re
from pathlib import Path
from typing import List, Dict, Any

# Dependencies check
try:
    import ezdxf
    import pdfplumber
    from docx import Document
except ImportError:
    print("⚠️  Missing dependencies. Run: pip install ezdxf pdfplumber python-docx")
    exit(1)

# ============================================================================
# CONFIGURATION
# ============================================================================

SOURCE_DIR = Path("project_documents")  # Directory containing PDFs, DXFs, and the Taxonomy DOCX
OUTPUT_FILE = Path("data/training_data/project_dialogue.jsonl")

# Fallback Contexts if Taxonomy file is missing
DEFAULT_SDIO_CONTEXT = {
    "100": "Project Requirement / Scope",
    "200": "Architectural Plan",
    "260": "Plumbing Plan",
    "261": "Plumbing Plan", 
    "270": "Mechanical Plan",
    "271": "Mechanical Plan",
    "274": "Mechanical Detail",
    "280": "Electrical Plan",
    "282": "Electrical Plan",
    "300": "Product Data",
    "302": "Product Data Submittal",
    "500": "Shop Drawing",
}

# ============================================================================
# EXTRACTORS
# ============================================================================

def load_taxonomy(source_dir: Path) -> Dict[str, str]:
    """Dynamically loads SDIO codes from the provided DOCX file."""
    taxonomy = DEFAULT_SDIO_CONTEXT.copy()
    
    # Find the taxonomy file
    docx_files = list(source_dir.glob("*Taxonomy*.docx"))
    if not docx_files:
        print("ℹ️  No Taxonomy DOCX found. Using default codes.")
        return taxonomy

    print(f"📖 Learning codes from: {docx_files[0].name}")
    try:
        doc = Document(docx_files[0])
        # Simple parser: looks for patterns like "271 - Mechanical Plan" or table rows
        for para in doc.paragraphs:
            text = para.text.strip()
            match = re.search(r"^(\d{3})\s*[-–]\s*(.+)", text)
            if match:
                code, desc = match.groups()
                taxonomy[code] = desc.strip()
                
        # Also check tables if the doc uses them
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2 and re.match(r"^\d{3}$", cells[0]):
                    taxonomy[cells[0]] = cells[1]
                    
    except Exception as e:
        print(f"⚠️  Error reading taxonomy: {e}")
        
    return taxonomy

def extract_dxf_text(file_path: Path) -> str:
    """Extracts text annotations from DXF drawings."""
    try:
        doc = ezdxf.readfile(file_path)
        msp = doc.modelspace()
        text_content = []
        for mtext in msp.query("MTEXT"):
            text_content.append(mtext.text)
        for text in msp.query("TEXT"):
            text_content.append(text.dxf.text)
        return "\n".join(text_content)
    except Exception as e:
        return f"[Error parsing DXF: {e}]"

def extract_pdf_text(file_path: Path) -> str:
    """Extracts text from PDF plans and specs."""
    text_content = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
        return "\n".join(text_content)
    except Exception as e:
        return f"[Error parsing PDF: {e}]"

# ============================================================================
# DIALOGUE GENERATOR
# ============================================================================

def generate_dialogue_pairs(filename: str, context_type: str, content: str) -> List[Dict]:
    pairs = []
    clean_content = re.sub(r'\s+', ' ', content).strip()[:3000] # Increased limit
    
    # 1. Identity Question
    pairs.append({
        "messages": [
            {"role": "user", "content": f"What is the file {filename}?"},
            {"role": "assistant", "content": f"This is the {context_type}. It contains the following technical data: {clean_content[:200]}..."}
        ]
    })
    
    # 2. Domain-Specific Questions
    if "Plumbing" in context_type or "26" in filename:
        pairs.append({
            "messages": [
                {"role": "user", "content": "What plumbing penetrations or drains are shown?"},
                {"role": "assistant", "content": f"The plumbing plan details the following fixtures and locations: {clean_content}"}
            ]
        })
    elif "Mechanical" in context_type or "27" in filename:
        pairs.append({
            "messages": [
                {"role": "user", "content": "Where are the RTUs or mechanical units located?"},
                {"role": "assistant", "content": f"The mechanical plan indicates unit locations and curbs here: {clean_content}"}
            ]
        })
    elif "Electrical" in context_type or "28" in filename:
        pairs.append({
            "messages": [
                {"role": "user", "content": "Are there any electrical conduits or feeds on the roof?"},
                {"role": "assistant", "content": f"The electrical plan shows the following roof-mounted equipment: {clean_content}"}
            ]
        })
    elif "Submittal" in context_type or "302" in filename:
        pairs.append({
            "messages": [
                {"role": "user", "content": "What are the product specifications for this submittal?"},
                {"role": "assistant", "content": f"This product data submittal specifies: {clean_content}"}
            ]
        })
    
    return pairs

# ============================================================================
# MAIN
# ============================================================================

def main():
    if not SOURCE_DIR.exists():
        SOURCE_DIR.mkdir(parents=True)
        print(f"⚠️ Created {SOURCE_DIR}. Place your files (PDF, DXF, DOCX) here and rerun.")
        return

    # 1. Load the Brain (Taxonomy)
    sdio_map = load_taxonomy(SOURCE_DIR)
    
    all_dialogue = []
    print(f"\n📂 Scanning {SOURCE_DIR}...")
    
    for file_path in SOURCE_DIR.glob("*.*"):
        if file_path.suffix.lower() not in ['.pdf', '.dxf']:
            continue
            
        print(f"Processing: {file_path.name}")
        
        # 2. Identify Context via SDIO Code
        # Looks for the first 3 digits in the filename (e.g., "271" in "271-IRO25...")
        code_match = re.match(r"^(\d{3})", file_path.name)
        if code_match:
            code = code_match.group(1)
            context = sdio_map.get(code, "Project Document")
        else:
            # Fallback for named files like "SDIO_Universal_Taxonomy"
            if "taxonomy" in file_path.name.lower(): continue 
            context = "General Document"

        # 3. Extract
        content = ""
        if file_path.suffix.lower() == '.dxf':
            content = extract_dxf_text(file_path)
        elif file_path.suffix.lower() == '.pdf':
            content = extract_pdf_text(file_path)
            
        if not content.strip():
            print(f"  ⚠️  Skipping empty file: {file_path.name}")
            continue

        # 4. Generate
        pairs = generate_dialogue_pairs(file_path.name, context, content)
        all_dialogue.extend(pairs)
        print(f"  ✅ Generated {len(pairs)} interactions for {context}")

    # 5. Save
    if OUTPUT_FILE.parent.exists() is False:
        OUTPUT_FILE.parent.mkdir(parents=True)
        
    with open(OUTPUT_FILE, 'w') as f:
        for pair in all_dialogue:
            f.write(json.dumps(pair) + '\n')
            
    print(f"\n🎉 Success! Saved {len(all_dialogue)} training conversations to {OUTPUT_FILE}")

if __name__ == "__main__":
    main()